from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import os

def create_document():
    document = Document()

    # Title
    title = document.add_heading('Face Authentication & Attendance System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle / Info
    p = document.add_paragraph('Detailed Technical Report')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_page_break()

    # 1. Project Overview
    document.add_heading('1. Project Overview', level=1)
    document.add_paragraph(
        "This project is a secure, biometric attendance system that uses Facial Recognition to authenticate users. "
        "It is designed to be robust against simple spoofing attacks (like holding up a photo) by implementing "
        "Liveness Detection via eye-blink analysis."
    )
    document.add_paragraph(
        "The system is built as a web application, allowing users to register their face and 'Punch In/Out' using any device with a camera."
    )

    document.add_heading('Key Features', level=2)
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Facial Recognition').bold = True
    p.add_run(': Uses the LBPH (Local Binary Patterns Histograms) algorithm for recognizing registered users.')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Liveness Detection').bold = True
    p.add_run(': Uses MediaPipe Face Mesh to detect eye blinks, ensuring the user is a real person.')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Web Interface').bold = True
    p.add_run(': A modern, responsive UI for easy interaction.')

    p = document.add_paragraph(style='List Bullet')
    p.add_run('JSON Database').bold = True
    p.add_run(': Lightweight, file-based storage for user data and attendance logs.')

    # 2. System Architecture
    document.add_heading('2. System Architecture', level=1)
    document.add_paragraph(
        "The system follows a standard Client-Server architecture. The Client (Browser) captures video frames, "
        "and the Server (Python/FastAPI) processes them using Computer Vision algorithms."
    )
    
    document.add_paragraph("[PLACEHOLDER: Insert Architecture Diagram Here]")
    document.add_paragraph("User -> UI -> Camera -> API -> Liveness Check -> Recognition -> Database")

    # 3. Technical Implementation Details
    document.add_heading('3. Technical Implementation Details', level=1)
    
    document.add_heading('3.1. Tech Stack', level=2)
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Language: ').bold = True
    p.add_run('Python 3.10+')
    
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Web Framework: ').bold = True
    p.add_run('FastAPI (High performance, async support)')

    p = document.add_paragraph(style='List Bullet')
    p.add_run('Computer Vision: ').bold = True
    p.add_run('OpenCV (Image processing, LBPH), MediaPipe (Face Landmarking)')

    document.add_heading('3.2. Liveness Detection (Anti-Spoofing)', level=2)
    document.add_paragraph("Goal: Prevent users from logging in using a photo of another person.")
    document.add_paragraph("Method: Blink Detection.")
    
    steps = [
        "The frontend captures a burst of 10 images over ~1.5 seconds.",
        "The backend calculates the Eye Aspect Ratio (EAR) for each frame.",
        "EAR is a geometric measurement of how 'open' the eye is.",
        "High EAR (> 0.22) = Open Eyes.",
        "Low EAR (< 0.18) = Closed Eyes (Blink).",
        "Verification: The system confirms liveness ONLY if it detects both an 'Open' state and a 'Closed' state within the image sequence."
    ]
    for step in steps:
        document.add_paragraph(step, style='List Number')

    document.add_heading('3.3. Face Recognition', level=2)
    document.add_paragraph("Method: LBPH (Local Binary Patterns Histograms).")
    document.add_paragraph(
        "LBPH is robust to local lighting changes. It divides the face into small cells, compares each pixel to its neighbors, "
        "and builds a histogram. During recognition, it compares the histogram of the input face to the training data."
    )

    # 4. Workflows & Usage
    document.add_heading('4. Workflows & Usage', level=1)

    document.add_heading('4.1. Registration Process', level=2)
    document.add_paragraph("1. User enters their name.")
    document.add_paragraph("2. Sits in front of the camera.")
    document.add_paragraph("3. System captures 10 training images.")
    document.add_paragraph("4. The model is immediately retrained to include the new user.")
    
    p = document.add_paragraph()
    runner = p.add_run("[INSERT SCREENSHOT HERE: Registration Page]")
    runner.font.color.rgb = RGBColor(255, 0, 0)
    runner.bold = True

    document.add_heading('4.2. Attendance "Punch In"', level=2)
    document.add_paragraph("1. User clicks 'Punch IN'.")
    document.add_paragraph("2. System prompts: 'Look at the camera and BLINK'.")
    document.add_paragraph("3. User blinks naturally. System captures the frame sequence.")
    document.add_paragraph("4. Backend verifies the blink, identifies the user, and logs the time.")

    p = document.add_paragraph()
    runner = p.add_run("[INSERT SCREENSHOT HERE: Punch In Success]")
    runner.font.color.rgb = RGBColor(255, 0, 0)
    runner.bold = True

    document.add_heading('4.3. Viewing History', level=2)
    document.add_paragraph("1. User navigates to the History page.")
    document.add_paragraph("2. A table displays all clock-in/out events with timestamps.")

    p = document.add_paragraph()
    runner = p.add_run("[INSERT SCREENSHOT HERE: History Log Table]")
    runner.font.color.rgb = RGBColor(255, 0, 0)
    runner.bold = True

    # 5. Project Structure
    document.add_heading('5. Project Structure Overview', level=1)
    structure = [
        ("main.py", "The entry point. Handles all web requests (API endpoints)."),
        ("core/auth.py", "The 'Brain'. Handles Liveness and LBPH Recognition."),
        ("core/db.py", "Handles reading/writing to data.json."),
        ("static/ & templates/", "Frontend User Interface."),
        ("data/", "Stores raw face images, trained model, and database.")
    ]
    
    for filename, desc in structure:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(filename).bold = True
        p.add_run(f": {desc}")

    # Save
    if not os.path.exists('docs'):
        os.makedirs('docs')
        
    output_path = os.path.join('docs', 'Face_Authentication_Report.docx')
    document.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_document()
